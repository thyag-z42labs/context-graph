# Copyright 2026 Neo4j Labs
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Word (``.docx``) parser for the Local File document connector.

Uses ``python-docx``. Heading detection is style-based: a paragraph whose
``style.name`` starts with ``"Heading "`` is treated as a heading; the
trailing digit gives the level (``"Heading 1"`` → 1, ``"Heading 3"`` →
3). Hyperlinks come from ``<w:hyperlink>`` elements in the paragraph's
XML, resolved against the document's relationships map.
"""

from __future__ import annotations

from pathlib import Path

from create_context_graph.connectors._local_file.parser import (
    ParsedDocument,
    ParsedSection,
    posix_uri,
)

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def parse(path: str | Path) -> ParsedDocument:
    """Parse a Word ``.docx`` file into a :class:`ParsedDocument`.

    Raises:
        ImportError: if ``python-docx`` is not installed.
    """
    try:
        import docx as _docx_lib
    except ImportError as exc:  # pragma: no cover - exercised at runtime only.
        raise ImportError(
            "Word parsing requires 'python-docx'. "
            "Install with: pip install 'python-docx>=1.2'"
        ) from exc

    p = Path(path)
    document = _docx_lib.Document(str(p))

    title = _document_title(document, p)

    # Walk the paragraphs once, classifying each as heading or body. For
    # every paragraph we also extract hyperlinks so we can attach them to
    # the right section.
    rels = document.part.rels
    paragraphs: list[dict] = []
    for para in document.paragraphs:
        style = (para.style.name if para.style else "") or ""
        level = _heading_level(style)
        text = para.text
        links = _paragraph_links(para, rels)
        paragraphs.append({
            "text": text,
            "level": level,
            "links": links,
        })

    preamble_lines: list[str] = []
    preamble_links: list[str] = []
    first_heading_idx = None
    for i, item in enumerate(paragraphs):
        if item["level"] is not None:
            first_heading_idx = i
            break
        preamble_lines.append(item["text"])
        for href in item["links"]:
            if href not in preamble_links:
                preamble_links.append(href)
    author, language = _docx_author_language(document)
    if first_heading_idx is None:
        return ParsedDocument(
            uri=posix_uri(p),
            title=title,
            preamble="\n".join(preamble_lines).strip(),
            sections=[],
            links=preamble_links,
            source_type="LOCAL_FILE",
            author=author,
            language=language,
        )

    sections = _build_section_tree(paragraphs[first_heading_idx:])
    return ParsedDocument(
        uri=posix_uri(p),
        title=title,
        preamble="\n".join(preamble_lines).strip(),
        sections=sections,
        links=preamble_links,
        source_type="LOCAL_FILE",
        author=author,
        language=language,
    )


# ---------------------------------------------------------------------------
# Internals
# ---------------------------------------------------------------------------


def _docx_author_language(document) -> tuple[str | None, str | None]:
    """Extract author and language from DOCX core properties."""
    author = language = None
    try:
        core = document.core_properties
        if core is not None:
            raw = getattr(core, "author", None)
            author = str(raw).strip() or None if raw else None
            raw = getattr(core, "language", None)
            language = str(raw).strip() or None if raw else None
    except Exception:  # pragma: no cover - missing or malformed core properties.
        pass
    return author, language


def _document_title(document, path: Path) -> str:
    """Use the document's core-properties title if present, else first H1,
    else filename stem."""
    try:
        core = document.core_properties
        if core is not None and core.title:
            return core.title.strip()
    except Exception:  # pragma: no cover - missing or malformed core properties.
        pass
    for para in document.paragraphs:
        style = (para.style.name if para.style else "") or ""
        if _heading_level(style) == 1 and para.text.strip():
            return para.text.strip()
    return path.stem


def _heading_level(style_name: str) -> int | None:
    """Return ``1..6`` for ``"Heading N"`` styles, else ``None``.

    ``"Title"`` is treated as level 1 (Word's default document-title
    style).
    """
    if not style_name:
        return None
    if style_name == "Title":
        return 1
    if style_name.startswith("Heading "):
        rest = style_name[len("Heading "):].strip()
        if rest.isdigit():
            n = int(rest)
            if 1 <= n <= 9:
                return min(n, 6)
    return None


def _paragraph_links(para, rels) -> list[str]:
    """Extract hyperlinks from a paragraph's XML.

    ``python-docx`` exposes the underlying ``lxml`` element via
    ``paragraph._element``. ``<w:hyperlink>`` elements carry an ``r:id``
    attribute that maps to a ``Relationship`` in ``rels``; its ``target``
    is the URL.
    """
    href_list: list[str] = []
    seen: set[str] = set()
    element = para._element
    for hyperlink in element.iter(f"{W_NS}hyperlink"):
        rid = hyperlink.get(f"{R_NS}id")
        if not rid:
            anchor = hyperlink.get(f"{W_NS}anchor")
            if anchor:
                href = "#" + anchor
                if href not in seen:
                    seen.add(href)
                    href_list.append(href)
            continue
        rel = rels.get(rid)
        if rel is None:
            continue
        target = getattr(rel, "target_ref", None) or getattr(rel, "target", None)
        if target and target not in seen:
            seen.add(str(target))
            href_list.append(str(target))
    return href_list


def _build_section_tree(items: list[dict]) -> list[ParsedSection]:
    """Build the nested ParsedSection tree from a flat heading/body list.

    Direct body of a section is the body-paragraphs that follow its
    heading up to (but not including) the next heading paragraph,
    regardless of that heading's level. Descendants live on their own
    sections.
    """
    # Identify heading indices.
    heading_idxs = [i for i, it in enumerate(items) if it["level"] is not None]
    root: list[ParsedSection] = []
    stack: list[ParsedSection] = []

    for h_pos, h_idx in enumerate(heading_idxs):
        head = items[h_idx]
        # Body runs from h_idx+1 to the next heading's index (exclusive).
        next_h_idx = (
            heading_idxs[h_pos + 1] if h_pos + 1 < len(heading_idxs) else len(items)
        )
        body_items = items[h_idx + 1 : next_h_idx]
        body_text = "\n".join(b["text"] for b in body_items).strip()
        body_links: list[str] = []
        for b in body_items:
            for href in b["links"]:
                if href not in body_links:
                    body_links.append(href)
        # Heading paragraphs may themselves contain hyperlinks (e.g. an H2
        # whose text is a link). Attach those to the section too.
        for href in head["links"]:
            if href not in body_links:
                body_links.append(href)

        section = ParsedSection(
            title=head["text"].strip(),
            level=head["level"],
            body=body_text,
            subsections=[],
            links=body_links,
        )
        while stack and stack[-1].level >= head["level"]:
            stack.pop()
        if stack:
            stack[-1].subsections.append(section)
        else:
            root.append(section)
        stack.append(section)
    return root
