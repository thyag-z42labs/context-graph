# Copyright 2026 Neo4j Labs
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Unit tests for the Local File document connector.

Mirrors the test list in ``scratch/doc-connector-requirements-v2.md`` §11.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from create_context_graph.connectors._local_file.parser import (
    ParsedDocument,
    ParsedSection,
    SUPPORTED_EXTENSIONS,
    parse_file,
    posix_uri,
    slugify,
)

# Optional-dep guards: skip whole modules whose libraries aren't installed.
pytest.importorskip("markdown_it")
pytest.importorskip("mdit_py_plugins")

from create_context_graph.connectors._local_file.parsers import markdown as md_parser


# ---------------------------------------------------------------------------
# slugify
# ---------------------------------------------------------------------------


class TestSlugify:
    def test_basic_lowercase(self):
        assert slugify("Hello World") == "hello-world"

    def test_punctuation_collapsed(self):
        # Spec example: ## OAuth 2.0 — Token Exchange!
        assert slugify("OAuth 2.0 — Token Exchange!") == "oauth-2-0-token-exchange"

    def test_unicode_emoji_dropped(self):
        # Emoji is dropped by NFKD + ascii-ignore.
        assert slugify("Hello 🌍 World") == "hello-world"

    def test_unicode_accents_normalized(self):
        # NFKD decomposes é → e + combining mark; ascii-ignore drops the mark.
        assert slugify("Café Résumé") == "cafe-resume"
        assert slugify("über cool") == "uber-cool"

    def test_cjk_dropped(self):
        # All non-ASCII codepoints are stripped, leaving nothing.
        assert slugify("你好世界") == ""

    def test_consecutive_separators_collapsed(self):
        assert slugify("foo --- bar") == "foo-bar"
        assert slugify("foo___bar") == "foo___bar"  # underscore is allowed

    def test_strips_leading_trailing_hyphens(self):
        assert slugify("---foo---") == "foo"

    def test_empty(self):
        assert slugify("") == ""
        assert slugify("---") == ""
        assert slugify("   ") == ""

    def test_digits_preserved(self):
        assert slugify("Section 3.14") == "section-3-14"


# ---------------------------------------------------------------------------
# posix_uri
# ---------------------------------------------------------------------------


class TestPosixUri:
    def test_uri_posix_normalized(self, tmp_path):
        # Even on macOS/Linux the URI should be POSIX-form. The real cross-OS
        # guarantee is that ``as_posix()`` is used; we assert no backslashes
        # appear and the URI is absolute.
        f = tmp_path / "guide.md"
        f.write_text("# x", encoding="utf-8")
        uri = posix_uri(f)
        assert "\\" not in uri
        assert uri.endswith("/guide.md")
        assert uri.startswith("/")

    def test_uri_pure_windows_path_via_purepath(self):
        # We can't actually resolve a Windows path on POSIX, but we can verify
        # that as_posix() canonicalises drive-letter style strings.
        from pathlib import PureWindowsPath

        win = PureWindowsPath("C:\\docs\\guide.md")
        # PureWindowsPath.as_posix() flips the slashes.
        assert win.as_posix() == "C:/docs/guide.md"


# ---------------------------------------------------------------------------
# parse_file dispatch
# ---------------------------------------------------------------------------


class TestParseFileDispatch:
    def test_supported_extensions_documented(self):
        # The dispatch table must cover the formats listed in the spec.
        for ext in (".md", ".mdx", ".markdown", ".pdf", ".html", ".htm",
                    ".adoc", ".asciidoc", ".asc", ".docx"):
            assert ext in SUPPORTED_EXTENSIONS

    def test_unsupported_extension_raises(self, tmp_path):
        f = tmp_path / "file.xyz"
        f.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported file extension"):
            parse_file(f)

    def test_missing_file_raises(self, tmp_path):
        with pytest.raises(FileNotFoundError):
            parse_file(tmp_path / "nope.md")

    def test_dispatch_routes_markdown(self, tmp_path):
        f = tmp_path / "x.md"
        f.write_text("# Title\nbody", encoding="utf-8")
        doc = parse_file(f)
        assert isinstance(doc, ParsedDocument)
        assert doc.title == "Title"


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------


class TestMarkdownParser:
    def test_basic_h1_title(self, tmp_path):
        f = tmp_path / "a.md"
        f.write_text("# Hello\n\nbody\n", encoding="utf-8")
        doc = md_parser.parse(f)
        assert doc.title == "Hello"
        assert doc.sections[0].title == "Hello"
        assert doc.sections[0].level == 1
        assert doc.sections[0].body == "body"

    def test_title_falls_back_to_filename(self, tmp_path):
        f = tmp_path / "no_h1.md"
        f.write_text("## sub\nbody\n", encoding="utf-8")
        doc = md_parser.parse(f)
        assert doc.title == "no_h1"

    def test_preamble_extracted(self, tmp_path):
        f = tmp_path / "p.md"
        f.write_text("Preamble text.\n\nMore preamble.\n\n# H1\nx\n", encoding="utf-8")
        doc = md_parser.parse(f)
        assert "Preamble text" in doc.preamble
        assert "More preamble" in doc.preamble

    def test_nested_sections(self, tmp_path):
        f = tmp_path / "n.md"
        f.write_text(
            "# H1\n\nbody1\n\n## H2\n\nbody2\n\n### H3\n\nbody3\n",
            encoding="utf-8",
        )
        doc = md_parser.parse(f)
        h1 = doc.sections[0]
        assert h1.level == 1 and h1.title == "H1"
        assert h1.body == "body1"
        h2 = h1.subsections[0]
        assert h2.level == 2 and h2.title == "H2"
        h3 = h2.subsections[0]
        assert h3.level == 3 and h3.title == "H3"

    def test_shallow_body(self, tmp_path):
        # body text should NOT include descendants — that lives on child nodes.
        f = tmp_path / "s.md"
        f.write_text(
            "# Top\n\ntop body\n\n## Child\n\nchild body\n",
            encoding="utf-8",
        )
        doc = md_parser.parse(f)
        assert doc.sections[0].body == "top body"
        assert "child body" not in doc.sections[0].body

    def test_skipped_heading_levels(self, tmp_path):
        f = tmp_path / "skip.md"
        f.write_text("# H1\n\n### H3\n\nbody\n", encoding="utf-8")
        doc = md_parser.parse(f)
        h1 = doc.sections[0]
        h3 = h1.subsections[0]
        assert h3.level == 3  # accurate level preserved
        assert h3.title == "H3"

    def test_inline_link_captured(self, tmp_path):
        f = tmp_path / "l.md"
        f.write_text(
            "# Title\n\nSee [the docs](https://example.com/docs) for more.\n",
            encoding="utf-8",
        )
        doc = md_parser.parse(f)
        assert doc.sections[0].links == ["https://example.com/docs"]

    def test_reference_link_captured(self, tmp_path):
        f = tmp_path / "r.md"
        f.write_text(
            "# Title\n\nSee [the docs][ref] for more.\n\n[ref]: https://example.com/r\n",
            encoding="utf-8",
        )
        doc = md_parser.parse(f)
        assert "https://example.com/r" in doc.sections[0].links

    def test_anchor_link_captured(self, tmp_path):
        f = tmp_path / "a.md"
        f.write_text("# Title\n\n[anchor](#install)\n", encoding="utf-8")
        doc = md_parser.parse(f)
        assert doc.sections[0].links == ["#install"]

    def test_links_are_section_scoped(self, tmp_path):
        f = tmp_path / "scope.md"
        f.write_text(
            "# H1\n\n[a](https://a.example)\n\n## H2\n\n[b](https://b.example)\n",
            encoding="utf-8",
        )
        doc = md_parser.parse(f)
        h1 = doc.sections[0]
        h2 = h1.subsections[0]
        assert h1.links == ["https://a.example"]
        assert h2.links == ["https://b.example"]

    def test_encoding_utf8_with_replace_fallback(self, tmp_path):
        # Write bytes that are not valid UTF-8.
        f = tmp_path / "mixed.md"
        f.write_bytes(b"# Title\n\ncaf\xe9 (latin-1)\n")
        doc = md_parser.parse(f)  # must not raise
        # body should contain a replacement char where the invalid byte was.
        assert doc.title == "Title"
        assert "�" in doc.sections[0].body

    def test_idempotent(self, tmp_path):
        f = tmp_path / "i.md"
        f.write_text(
            "# A\n\nbody-a\n\n## B\n\nbody-b\n",
            encoding="utf-8",
        )
        d1 = md_parser.parse(f)
        d2 = md_parser.parse(f)
        assert d1 == d2


# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------


def _make_pdf_with_outline(path: Path) -> None:
    """Generate a tiny multi-page PDF with a nested outline.

    Returns nothing; writes to ``path``.
    """
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    c = canvas.Canvas(str(path), pagesize=letter)
    c.setTitle("Test Document")

    # Page 1 — Chapter 1 H1
    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, 720, "Chapter 1")
    c.setFont("Helvetica", 12)
    c.drawString(72, 690, "Body text for chapter 1.")
    # Add outline entry pointing at this page.
    c.bookmarkPage("ch1")
    c.addOutlineEntry("Chapter 1", "ch1", level=0)
    c.showPage()

    # Page 2 — Section 1.1 H2 (level=1)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72, 720, "Section 1.1")
    c.setFont("Helvetica", 12)
    c.drawString(72, 690, "Section body text.")
    # External hyperlink annotation.
    c.linkURL("https://example.com/x", (72, 660, 200, 680))
    c.drawString(72, 660, "see example")
    c.bookmarkPage("s11")
    c.addOutlineEntry("Section 1.1", "s11", level=1)
    c.showPage()

    # Page 3 — Chapter 2 H1
    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, 720, "Chapter 2")
    c.setFont("Helvetica", 12)
    c.drawString(72, 690, "Chapter 2 body text.")
    c.bookmarkPage("ch2")
    c.addOutlineEntry("Chapter 2", "ch2", level=0)
    c.showPage()

    c.save()


def _make_pdf_no_outline(path: Path) -> None:
    """Generate a PDF with two visually distinct font sizes but no outline."""
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    c = canvas.Canvas(str(path), pagesize=letter)
    c.setTitle("Heuristic Doc")
    # Large heading-sized line.
    c.setFont("Helvetica-Bold", 24)
    c.drawString(72, 720, "Big Heading")
    # Body-sized lines — must be the dominant size by character count.
    c.setFont("Helvetica", 10)
    for i in range(20):
        c.drawString(72, 700 - i * 12, "Body line " + "x" * 40)
    c.setFont("Helvetica-Bold", 24)
    c.drawString(72, 380, "Another Big Heading")
    c.setFont("Helvetica", 10)
    for i in range(20):
        c.drawString(72, 360 - i * 12, "More body line " + "y" * 40)
    c.showPage()
    c.save()


class TestPDFParser:
    def test_fetch_pdf_with_outline(self, tmp_path):
        pytest.importorskip("reportlab")
        pdf_path = tmp_path / "outline.pdf"
        _make_pdf_with_outline(pdf_path)
        from create_context_graph.connectors._local_file.parsers import (
            pdf as pdf_parser,
        )
        doc = pdf_parser.parse(pdf_path)
        # Title from metadata.
        assert doc.title == "Test Document"
        # Two top-level sections (Chapter 1, Chapter 2).
        titles = [s.title for s in doc.sections]
        assert "Chapter 1" in titles
        assert "Chapter 2" in titles
        # Chapter 1 should have Section 1.1 as a level=2 child.
        ch1 = next(s for s in doc.sections if s.title == "Chapter 1")
        assert any(ss.title == "Section 1.1" and ss.level == 2 for ss in ch1.subsections)
        # External URI link annotation collected.
        assert "https://example.com/x" in doc.links

    def test_fetch_pdf_outline_deterministic(self, tmp_path):
        pytest.importorskip("reportlab")
        pdf_path = tmp_path / "det.pdf"
        _make_pdf_with_outline(pdf_path)
        from create_context_graph.connectors._local_file.parsers import (
            pdf as pdf_parser,
        )
        d1 = pdf_parser.parse(pdf_path)
        d2 = pdf_parser.parse(pdf_path)
        # Same URIs, same section structure, same titles, same levels.
        assert [s.title for s in d1.sections] == [s.title for s in d2.sections]
        assert [s.level for s in d1.sections] == [s.level for s in d2.sections]
        for s1, s2 in zip(d1.sections, d2.sections):
            assert [ss.title for ss in s1.subsections] == [ss.title for ss in s2.subsections]

    def test_fetch_pdf_without_outline_fallback_to_heuristics(self, tmp_path):
        pytest.importorskip("reportlab")
        pdf_path = tmp_path / "noopl.pdf"
        _make_pdf_no_outline(pdf_path)
        from create_context_graph.connectors._local_file.parsers import (
            pdf as pdf_parser,
        )
        doc = pdf_parser.parse(pdf_path)
        # Title from metadata.
        assert doc.title == "Heuristic Doc"
        # We expect at least one heading detected (largest font is "Big Heading").
        assert len(doc.sections) >= 1
        # The largest font should map to level 1.
        assert any(s.level == 1 and "Big Heading" in s.title for s in doc.sections)


class TestPDFParserPdfOxide:
    """Tier 0 (pdf-oxide) tests — pdf-oxide is bundled so no mock needed for happy path."""

    def test_tier0_outline_pdf(self):
        """pdf-oxide should parse the fixture PDF and produce sections from the outline."""
        pytest.importorskip("pdf_oxide")
        from create_context_graph.connectors._local_file.parsers import pdf as pdf_parser

        fixture = Path("tests/fixtures/local_file_vault/external/acme-10q-2026q1.pdf")
        if not fixture.exists():
            pytest.skip("fixture PDF not found")

        doc = pdf_parser.parse(fixture)
        assert doc.title  # should have a title (from first H1)
        assert doc.sections  # outline-bearing PDF should have sections
        top_titles = [s.title for s in doc.sections]
        assert any("Part I" in t or "ACME" in t or "Financial" in t for t in top_titles)

    def test_tier0_heading_structure(self):
        """Sections parsed from the fixture PDF should have nested subsections."""
        pytest.importorskip("pdf_oxide")
        from create_context_graph.connectors._local_file.parsers import pdf as pdf_parser

        fixture = Path("tests/fixtures/local_file_vault/external/acme-10q-2026q1.pdf")
        if not fixture.exists():
            pytest.skip("fixture PDF not found")

        doc = pdf_parser.parse(fixture)
        all_sections = []

        def collect(sections):
            for s in sections:
                all_sections.append(s)
                collect(s.subsections)

        collect(doc.sections)
        assert len(all_sections) >= 3

    def test_tier0_markdown_parse_sections(self):
        """_parse_markdown_sections should build correct hierarchy from markdown."""
        from create_context_graph.connectors._local_file.parsers.pdf import _parse_markdown_sections

        md = "# Title\n\nIntro text.\n\n## Section A\n\nBody A.\n\n### Sub A1\n\nSub body.\n\n## Section B\n\nBody B."
        sections, preamble = _parse_markdown_sections(md)

        assert preamble == ""
        assert len(sections) == 1  # "Title" is the single root
        title_sec = sections[0]
        assert title_sec.title == "Title"
        assert len(title_sec.subsections) == 2
        assert title_sec.subsections[0].title == "Section A"
        assert len(title_sec.subsections[0].subsections) == 1
        assert title_sec.subsections[0].subsections[0].title == "Sub A1"

    def test_tier0_title_from_first_h1(self):
        """_title_from_markdown should use the first H1 line."""
        from create_context_graph.connectors._local_file.parsers.pdf import _title_from_markdown

        assert _title_from_markdown("# My Doc\n\nsome text", "fallback") == "My Doc"
        assert _title_from_markdown("no heading here", "fallback") == "fallback"
        assert _title_from_markdown("", "stem") == "stem"

    def test_tier0_uri_links_extracted(self):
        """_pdf_oxide_uri_links should collect inline and bare URLs."""
        from create_context_graph.connectors._local_file.parsers.pdf import _pdf_oxide_uri_links

        md = "See [docs](https://example.com/docs) and https://other.org/page for details."
        links = _pdf_oxide_uri_links(md)
        assert "https://example.com/docs" in links
        assert "https://other.org/page" in links

    def test_tier0_falls_through_on_import_error(self, tmp_path):
        """When pdf_oxide is absent, parse() should fall through to pypdf tiers."""
        pytest.importorskip("reportlab")
        import sys
        import unittest.mock as mock
        from create_context_graph.connectors._local_file.parsers import pdf as pdf_parser

        pdf_path = tmp_path / "nopdfoxide.pdf"
        _make_pdf_with_outline(pdf_path)

        with mock.patch.dict(sys.modules, {"pdf_oxide": None}):
            doc = pdf_parser.parse(pdf_path)

        assert doc.title == "Test Document"
        assert any(s.title == "Chapter 1" for s in doc.sections)


# ---------------------------------------------------------------------------
# HTML parser
# ---------------------------------------------------------------------------


# Optional-dep guard for HTML.
pytest.importorskip("bs4")
from create_context_graph.connectors._local_file.parsers import html as html_parser  # noqa: E402


class TestHTMLParser:
    def test_basic_h1_title(self, tmp_path):
        f = tmp_path / "h.html"
        f.write_text(
            "<html><body><h1>Hello</h1><p>body</p></body></html>",
            encoding="utf-8",
        )
        doc = html_parser.parse(f)
        assert doc.title == "Hello"
        assert doc.sections[0].title == "Hello"
        assert doc.sections[0].level == 1
        assert "body" in doc.sections[0].body

    def test_title_tag_when_no_h1(self, tmp_path):
        f = tmp_path / "h.html"
        f.write_text(
            "<html><head><title>Doc Title</title></head>"
            "<body><h2>only h2</h2></body></html>",
            encoding="utf-8",
        )
        doc = html_parser.parse(f)
        assert doc.title == "Doc Title"

    def test_filename_fallback(self, tmp_path):
        f = tmp_path / "stem_doc.html"
        f.write_text("<html><body><p>no headings</p></body></html>", encoding="utf-8")
        doc = html_parser.parse(f)
        assert doc.title == "stem_doc"

    def test_nested_sections(self, tmp_path):
        f = tmp_path / "n.html"
        f.write_text(
            "<html><body>"
            "<h1>A</h1><p>a body</p>"
            "<h2>B</h2><p>b body</p>"
            "<h3>C</h3><p>c body</p>"
            "</body></html>",
            encoding="utf-8",
        )
        doc = html_parser.parse(f)
        a = doc.sections[0]
        assert a.title == "A" and a.level == 1
        b = a.subsections[0]
        assert b.title == "B" and b.level == 2
        c = b.subsections[0]
        assert c.title == "C" and c.level == 3

    def test_shallow_body(self, tmp_path):
        f = tmp_path / "s.html"
        f.write_text(
            "<html><body><h1>Top</h1><p>top body</p>"
            "<h2>Child</h2><p>child body</p></body></html>",
            encoding="utf-8",
        )
        doc = html_parser.parse(f)
        top = doc.sections[0]
        assert "top body" in top.body
        assert "child body" not in top.body

    def test_links_captured(self, tmp_path):
        f = tmp_path / "l.html"
        f.write_text(
            "<html><body><h1>T</h1>"
            "<p><a href='https://example.com/x'>x</a></p>"
            "<p><a href='#anchor'>a</a></p>"
            "</body></html>",
            encoding="utf-8",
        )
        doc = html_parser.parse(f)
        links = doc.sections[0].links
        assert "https://example.com/x" in links
        assert "#anchor" in links

    def test_img_not_a_link(self, tmp_path):
        f = tmp_path / "i.html"
        f.write_text(
            "<html><body><h1>T</h1>"
            "<p><img src='https://example.com/pic.png' /></p>"
            "</body></html>",
            encoding="utf-8",
        )
        doc = html_parser.parse(f)
        assert doc.sections[0].links == []

    def test_skipped_heading_levels(self, tmp_path):
        f = tmp_path / "sk.html"
        f.write_text(
            "<html><body><h1>A</h1><p>x</p><h3>C</h3><p>y</p></body></html>",
            encoding="utf-8",
        )
        doc = html_parser.parse(f)
        a = doc.sections[0]
        c = a.subsections[0]
        assert c.level == 3 and c.title == "C"

    def test_no_headings_returns_preamble(self, tmp_path):
        f = tmp_path / "n.html"
        f.write_text("<html><body><p>just text</p></body></html>", encoding="utf-8")
        doc = html_parser.parse(f)
        assert doc.sections == []
        assert "just text" in doc.preamble

    def test_idempotent(self, tmp_path):
        f = tmp_path / "i.html"
        f.write_text(
            "<html><body><h1>A</h1><p>x</p><h2>B</h2><p>y</p></body></html>",
            encoding="utf-8",
        )
        d1 = html_parser.parse(f)
        d2 = html_parser.parse(f)
        assert d1 == d2


# ---------------------------------------------------------------------------
# AsciiDoc parser
# ---------------------------------------------------------------------------


from create_context_graph.connectors._local_file.parsers import asciidoc as adoc_parser  # noqa: E402


class TestAsciiDocParser:
    def test_basic_heading_levels(self, tmp_path):
        f = tmp_path / "a.adoc"
        f.write_text(
            "= Top\n\nbody\n\n== Sub A\n\nsub a body\n\n=== Deep\n\ndeep body\n",
            encoding="utf-8",
        )
        doc = adoc_parser.parse(f)
        assert doc.title == "Top"
        top = doc.sections[0]
        assert top.level == 1
        assert top.subsections[0].level == 2
        assert top.subsections[0].subsections[0].level == 3

    def test_title_filename_fallback(self, tmp_path):
        f = tmp_path / "no_h1.adoc"
        f.write_text("== H2 Only\nbody\n", encoding="utf-8")
        doc = adoc_parser.parse(f)
        assert doc.title == "no_h1"

    def test_skip_literal_block(self, tmp_path):
        # '=' inside ---- (listing) block must NOT be treated as a heading.
        f = tmp_path / "lit.adoc"
        f.write_text(
            "= Real H1\n\nbody\n\n----\n= NOT A HEADING\n=== ALSO NOT A HEADING\n----\n",
            encoding="utf-8",
        )
        doc = adoc_parser.parse(f)
        # Only one real heading should be found.
        assert len(doc.sections) == 1
        assert doc.sections[0].title == "Real H1"
        # No nested subsections from inside the block.
        assert doc.sections[0].subsections == []

    def test_skip_passthrough_and_table_blocks(self, tmp_path):
        f = tmp_path / "p.adoc"
        f.write_text(
            "= H1\n"
            "\n"
            "++++\n= passthrough not heading ++++\n++++\n"
            "\n"
            "|===\n| col1 | = inside table not heading\n|===\n",
            encoding="utf-8",
        )
        doc = adoc_parser.parse(f)
        assert len(doc.sections) == 1
        assert doc.sections[0].title == "H1"

    def test_link_macro_captured(self, tmp_path):
        f = tmp_path / "lm.adoc"
        f.write_text(
            "= T\n\nSee link:https://example.com/docs[the docs].\n",
            encoding="utf-8",
        )
        doc = adoc_parser.parse(f)
        assert "https://example.com/docs" in doc.sections[0].links

    def test_autolink_captured(self, tmp_path):
        f = tmp_path / "al.adoc"
        f.write_text(
            "= T\n\nSee https://example.com/page.html for info.\n",
            encoding="utf-8",
        )
        doc = adoc_parser.parse(f)
        # Trailing period should be stripped, not part of the URL.
        assert "https://example.com/page.html" in doc.sections[0].links

    def test_xref_captured(self, tmp_path):
        f = tmp_path / "x.adoc"
        f.write_text("= T\n\nSee <<my-anchor,more>>.\n", encoding="utf-8")
        doc = adoc_parser.parse(f)
        assert "#my-anchor" in doc.sections[0].links

    def test_idempotent(self, tmp_path):
        f = tmp_path / "i.adoc"
        f.write_text("= A\n\nbody\n\n== B\n\nb body\n", encoding="utf-8")
        d1 = adoc_parser.parse(f)
        d2 = adoc_parser.parse(f)
        assert d1 == d2


# ---------------------------------------------------------------------------
# Word (.docx) parser
# ---------------------------------------------------------------------------


pytest.importorskip("docx")
from create_context_graph.connectors._local_file.parsers import docx as docx_parser  # noqa: E402


def _make_docx(path: Path, paragraphs: list[tuple[str, str]]) -> None:
    """Create a docx file from ``paragraphs`` — a list of ``(style, text)``."""
    pytest.importorskip("docx")
    import docx as _docx_lib

    doc = _docx_lib.Document()
    for style, text in paragraphs:
        doc.add_paragraph(text, style=style)
    doc.save(str(path))


class TestDocxParser:
    def test_basic_heading_levels(self, tmp_path):
        f = tmp_path / "a.docx"
        _make_docx(
            f,
            [
                ("Heading 1", "Top"),
                ("Normal", "Top body."),
                ("Heading 2", "Sub"),
                ("Normal", "Sub body."),
                ("Heading 3", "Deep"),
                ("Normal", "Deep body."),
            ],
        )
        doc = docx_parser.parse(f)
        assert doc.title == "Top"
        top = doc.sections[0]
        assert top.level == 1
        assert top.body == "Top body."
        sub = top.subsections[0]
        assert sub.level == 2
        deep = sub.subsections[0]
        assert deep.level == 3
        assert deep.title == "Deep"

    def test_shallow_body(self, tmp_path):
        f = tmp_path / "s.docx"
        _make_docx(
            f,
            [
                ("Heading 1", "Top"),
                ("Normal", "top body"),
                ("Heading 2", "Child"),
                ("Normal", "child body"),
            ],
        )
        doc = docx_parser.parse(f)
        top = doc.sections[0]
        assert "top body" in top.body
        assert "child body" not in top.body

    def test_no_heading_uses_filename(self, tmp_path):
        f = tmp_path / "stem.docx"
        _make_docx(f, [("Normal", "just body text")])
        doc = docx_parser.parse(f)
        # Title comes from filename stem.
        assert doc.title == "stem"
        assert "just body text" in doc.preamble

    def test_idempotent(self, tmp_path):
        f = tmp_path / "i.docx"
        _make_docx(
            f,
            [
                ("Heading 1", "A"),
                ("Normal", "x"),
                ("Heading 2", "B"),
                ("Normal", "y"),
            ],
        )
        d1 = docx_parser.parse(f)
        d2 = docx_parser.parse(f)
        assert d1 == d2


# ---------------------------------------------------------------------------
# DocumentMapper (ParsedDocument → NormalizedData)
# ---------------------------------------------------------------------------


from create_context_graph.connectors._local_file.mapper import DocumentMapper, map_documents  # noqa: E402


def _sec(title: str, level: int, body: str = "", links=None, subs=None) -> ParsedSection:
    return ParsedSection(
        title=title,
        level=level,
        body=body,
        links=list(links or []),
        subsections=list(subs or []),
    )


def _doc(uri: str, title: str, sections=None, preamble: str = "", links=None) -> ParsedDocument:
    return ParsedDocument(
        uri=uri,
        title=title,
        preamble=preamble,
        sections=list(sections or []),
        links=list(links or []),
        source_type="LOCAL_FILE",
    )


class TestDocumentMapper:
    def test_document_uri_is_name_field(self):
        d = _doc("/docs/x.md", "X")
        data = map_documents([d])
        assert data.entities["Document"][0]["name"] == "/docs/x.md"

    def test_section_uri_format(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[_sec("Installation", 2)],
        )
        data = map_documents([d])
        section = data.entities["Section"][0]
        assert section["name"] == "/docs/x.md#installation"
        assert section["title"] == "Installation"
        assert section["headingLevel"] == 2

    def test_nested_section_uri_uses_slash(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[
                _sec("Top", 1, subs=[
                    _sec("Sub", 2, subs=[
                        _sec("Deep", 3),
                    ]),
                ]),
            ],
        )
        data = map_documents([d])
        names = {s["name"] for s in data.entities["Section"]}
        assert "/docs/x.md#top" in names
        assert "/docs/x.md#top/sub" in names
        assert "/docs/x.md#top/sub/deep" in names

    def test_description_shallow_with_child_pointers(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[
                _sec("Install", 2, body="install body",
                     subs=[_sec("Python", 3), _sec("CLI", 3)]),
            ],
        )
        data = map_documents([d])
        install = next(
            s for s in data.entities["Section"] if s["name"] == "/docs/x.md#install"
        )
        assert "install body" in install["description"]
        assert "uri:/docs/x.md#install/python" in install["description"]
        assert "uri:/docs/x.md#install/cli" in install["description"]
        # Child body must NOT bleed into parent description.
        py = next(
            s for s in data.entities["Section"]
            if s["name"] == "/docs/x.md#install/python"
        )
        py["description"]  # exists
        # Document description references its top-level section.
        doc_node = data.entities["Document"][0]
        assert "uri:/docs/x.md#install" in doc_node["description"]

    def test_duplicate_headings_disambiguated(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[
                _sec("Install", 2),
                _sec("Install", 2),
                _sec("Install", 2),
            ],
        )
        data = map_documents([d])
        names = sorted(s["name"] for s in data.entities["Section"])
        assert names == [
            "/docs/x.md#install",
            "/docs/x.md#install-1",
            "/docs/x.md#install-2",
        ]

    def test_duplicate_headings_scoped_per_parent(self):
        # Two parents each with an "Overview" child — must NOT collide.
        d = _doc(
            "/docs/x.md", "X",
            sections=[
                _sec("A", 2, subs=[_sec("Overview", 3)]),
                _sec("B", 2, subs=[_sec("Overview", 3)]),
            ],
        )
        data = map_documents([d])
        names = {s["name"] for s in data.entities["Section"]}
        assert "/docs/x.md#a/overview" in names
        assert "/docs/x.md#b/overview" in names
        # Neither overview should have a "-1" disambiguator.
        assert "/docs/x.md#a/overview-1" not in names
        assert "/docs/x.md#b/overview-1" not in names

    def test_has_section_edges(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[
                _sec("Top", 1, subs=[_sec("Sub", 2)]),
            ],
        )
        data = map_documents([d])
        rels = data.relationships
        has_section = [r for r in rels if r["type"] == "HAS_SECTION"]
        # Doc → Top
        assert any(
            r["source_name"] == "/docs/x.md" and r["source_label"] == "Document"
            and r["target_name"] == "/docs/x.md#top"
            for r in has_section
        )
        # Top → Sub
        assert any(
            r["source_name"] == "/docs/x.md#top" and r["source_label"] == "Section"
            and r["target_name"] == "/docs/x.md#top/sub"
            for r in has_section
        )

    def test_skipped_heading_levels(self):
        # ParsedSection trees from the parsers already place H3 under H1
        # when H2 is skipped — the mapper just preserves that.
        d = _doc(
            "/docs/x.md", "X",
            sections=[
                _sec("Top", 1, subs=[_sec("Deep", 3)]),
            ],
        )
        data = map_documents([d])
        deep = next(
            s for s in data.entities["Section"]
            if s["name"] == "/docs/x.md#top/deep"
        )
        assert deep["headingLevel"] == 3

    def test_links_to_external_url(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[
                _sec("A", 2, links=["https://example.com/y"]),
            ],
        )
        data = map_documents([d])
        # Stub URL_LINK document created.
        url_stub = next(
            doc for doc in data.entities["Document"]
            if doc["name"] == "https://example.com/y"
        )
        assert url_stub["sourceType"] == "URL_LINK"
        assert "title" not in url_stub  # stub has no title
        # LINKS_TO edge present.
        assert any(
            r["type"] == "LINKS_TO"
            and r["source_name"] == "/docs/x.md#a"
            and r["target_name"] == "https://example.com/y"
            and r["target_label"] == "Document"
            for r in data.relationships
        )

    def test_links_to_relative_local_path(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[_sec("A", 2, links=["./api.md"])],
        )
        data = map_documents([d])
        assert any(
            doc["name"] == "/docs/api.md" and doc["sourceType"] == "LOCAL_FILE"
            for doc in data.entities["Document"]
        )

    def test_links_to_parent_relative_local_path(self):
        d = _doc(
            "/docs/sub/x.md", "X",
            sections=[_sec("A", 2, links=["../top.md"])],
        )
        data = map_documents([d])
        assert any(
            doc["name"] == "/docs/top.md"
            for doc in data.entities["Document"]
        )

    def test_links_to_absolute_local_path(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[_sec("A", 2, links=["/elsewhere/y.md"])],
        )
        data = map_documents([d])
        assert any(
            doc["name"] == "/elsewhere/y.md"
            for doc in data.entities["Document"]
        )

    def test_links_to_same_doc_anchor(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[_sec("A", 2, links=["#advanced"])],
        )
        data = map_documents([d])
        assert any(
            r["type"] == "LINKS_TO"
            and r["target_name"] == "/docs/x.md#advanced"
            and r["target_label"] == "Section"
            for r in data.relationships
        )

    def test_links_to_cross_doc_anchor(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[_sec("A", 2, links=["./other.md#install"])],
        )
        data = map_documents([d])
        assert any(
            r["type"] == "LINKS_TO"
            and r["target_name"] == "/docs/other.md#install"
            and r["target_label"] == "Section"
            for r in data.relationships
        )

    @pytest.mark.parametrize("href", [
        "mailto:noreply@example.com",
        "tel:+12025550100",
        "javascript:alert(1)",
        "data:text/plain;base64,aGVsbG8=",
        "ftp://example.com/file",
    ])
    def test_links_skip_non_document_schemes(self, href):
        d = _doc("/docs/x.md", "X", sections=[_sec("A", 2, links=[href])])
        data = map_documents([d])
        # No LINKS_TO for these.
        assert not any(r["type"] == "LINKS_TO" for r in data.relationships)
        # Only the source document exists in entities (no stub).
        assert len(data.entities.get("Document", [])) == 1

    def test_links_dedupe_across_sections(self):
        # Same URL referenced from 3 sections → 1 Document node, 3 LINKS_TO edges.
        d = _doc(
            "/docs/x.md", "X",
            sections=[
                _sec("A", 2, links=["https://shared.example/x"]),
                _sec("B", 2, links=["https://shared.example/x"]),
                _sec("C", 2, links=["https://shared.example/x"]),
            ],
        )
        data = map_documents([d])
        stubs = [
            doc for doc in data.entities["Document"]
            if doc["name"] == "https://shared.example/x"
        ]
        assert len(stubs) == 1
        edges = [
            r for r in data.relationships
            if r["type"] == "LINKS_TO" and r["target_name"] == "https://shared.example/x"
        ]
        assert len(edges) == 3

    def test_documents_list_empty(self):
        # Spec §5: do NOT populate the `documents=[]` list — avoids
        # collision with ingest.py's MERGE-on-title :Document nodes.
        d = _doc("/docs/x.md", "X", sections=[_sec("A", 2)])
        data = map_documents([d])
        assert data.documents == []

    def test_traces_list_empty(self):
        d = _doc("/docs/x.md", "X", sections=[_sec("A", 2)])
        data = map_documents([d])
        assert data.traces == []

    def test_stub_upgrade_when_target_in_run(self):
        # Ingest document Y first; then ingest X that links to Y; Y must NOT
        # be turned into a stub (it's already real).
        target = _doc("/docs/y.md", "Y", sections=[_sec("First", 2, body="b")])
        source = _doc(
            "/docs/x.md", "X",
            sections=[_sec("A", 2, links=["./y.md"])],
        )
        mapper = DocumentMapper()
        mapper.add(target)
        mapper.add(source)
        data = mapper.build()
        y_docs = [d for d in data.entities["Document"] if d["name"] == "/docs/y.md"]
        assert len(y_docs) == 1
        # Y retains its title (was added as a real doc).
        assert y_docs[0]["title"] == "Y"

    def test_reverse_order_link_resolution(self):
        # Ingest source first (creates stub for Y), then ingest Y (upgrades
        # stub) — final node should be the rich Y with title/description.
        source = _doc(
            "/docs/x.md", "X",
            sections=[_sec("A", 2, links=["./y.md"])],
        )
        target = _doc("/docs/y.md", "Y", sections=[_sec("First", 2, body="b")])
        mapper = DocumentMapper()
        mapper.add(source)
        mapper.add(target)
        data = mapper.build()
        y_docs = [d for d in data.entities["Document"] if d["name"] == "/docs/y.md"]
        assert len(y_docs) == 1
        assert y_docs[0]["title"] == "Y"  # upgraded
        assert y_docs[0].get("description")  # has content

    def test_anchor_link_creates_parent_document_stub(self):
        # A link to an anchor on an un-ingested document should create:
        # 1. A Section stub for the anchor target.
        # 2. A Document stub for the parent document.
        # 3. A HAS_SECTION edge between them so the graph is traversable.
        source = _doc(
            "/docs/x.md", "X",
            sections=[_sec("A", 2, links=["./y.md#intro"])],
        )
        mapper = DocumentMapper()
        mapper.add(source)
        data = mapper.build()

        section_stubs = [s for s in data.entities["Section"] if "#intro" in s["name"]]
        assert len(section_stubs) == 1, "Section stub for anchor expected"

        doc_stubs = [d for d in data.entities["Document"] if d["name"].endswith("/docs/y.md")]
        assert len(doc_stubs) == 1, "Parent Document stub expected"

        has_section = [
            r for r in data.relationships
            if r["type"] == "HAS_SECTION"
            and r["source_name"].endswith("/docs/y.md")
            and "#intro" in r["target_name"]
        ]
        assert len(has_section) == 1, "HAS_SECTION from stub Document to stub Section expected"

    def test_anchor_link_stub_upgraded_on_later_ingest(self):
        # After a stub Document + Section are created from an anchor link,
        # ingesting the real document should upgrade both in-place without
        # producing duplicates.
        source = _doc(
            "/docs/x.md", "X",
            sections=[_sec("A", 2, links=["./y.md#intro"])],
        )
        target = _doc(
            "/docs/y.md", "Y",
            sections=[_sec("Intro", 1, body="real body")],
        )
        mapper = DocumentMapper()
        mapper.add(source)
        mapper.add(target)
        data = mapper.build()

        # Exactly one Document node for y.md (stub upgraded, not duplicated).
        y_docs = [d for d in data.entities["Document"] if d["name"].endswith("/docs/y.md")]
        assert len(y_docs) == 1
        assert y_docs[0].get("title") == "Y"  # upgraded with real title

        # Exactly one HAS_SECTION from y.md to its intro section.
        has_section = [
            r for r in data.relationships
            if r["type"] == "HAS_SECTION" and r["source_name"].endswith("/docs/y.md")
        ]
        assert len(has_section) == 1

    def test_idempotent_mapping(self):
        d = _doc(
            "/docs/x.md", "X",
            sections=[_sec("A", 2, body="ab")],
        )
        a = map_documents([d])
        b = map_documents([d])
        # loadedAt differs across runs, but everything else should match.
        for entity_list in a.entities.values():
            for entity in entity_list:
                entity.pop("loadedAt", None)
        for entity_list in b.entities.values():
            for entity in entity_list:
                entity.pop("loadedAt", None)
        assert a.entities == b.entities
        assert a.relationships == b.relationships


# ---------------------------------------------------------------------------
# LocalFileConnector — registration, authenticate, fetch
# ---------------------------------------------------------------------------


from create_context_graph.connectors import (  # noqa: E402
    CONNECTOR_REGISTRY,
    NormalizedData,
    get_connector,
    merge_connector_results,
)
from create_context_graph.connectors.local_file_connector import LocalFileConnector  # noqa: E402


class TestLocalFileConnectorRegistration:
    def test_registration(self):
        assert "local-file" in CONNECTOR_REGISTRY
        assert CONNECTOR_REGISTRY["local-file"] is LocalFileConnector

    def test_service_metadata(self):
        conn = get_connector("local-file")
        assert conn.service_name == "Local File"
        assert "filesystem" in conn.service_description.lower()
        assert conn.requires_oauth is False

    def test_credential_prompts_returns_one_required_path_prompt(self):
        conn = get_connector("local-file")
        prompts = conn.get_credential_prompts()
        assert len(prompts) == 1
        assert prompts[0]["name"] == "paths"
        assert prompts[0]["secret"] is False


class TestLocalFileConnectorAuthenticate:
    def test_authenticate_with_list_of_paths(self, tmp_path):
        f = tmp_path / "a.md"
        f.write_text("# A", encoding="utf-8")
        conn = LocalFileConnector()
        conn.authenticate({"paths": [str(f)]})
        assert conn._paths == [f]

    def test_authenticate_with_comma_separated_string(self, tmp_path):
        a = tmp_path / "a.md"
        b = tmp_path / "b.md"
        a.write_text("# A", encoding="utf-8")
        b.write_text("# B", encoding="utf-8")
        conn = LocalFileConnector()
        conn.authenticate({"paths": f"{a}, {b}"})
        assert sorted(conn._paths) == sorted([a, b])

    def test_authenticate_empty_paths_raises(self):
        conn = LocalFileConnector()
        with pytest.raises(ValueError, match="local-file-path"):
            conn.authenticate({})

    def test_authenticate_empty_string_raises(self):
        conn = LocalFileConnector()
        with pytest.raises(ValueError):
            conn.authenticate({"paths": ""})

    def test_authenticate_carries_options(self, tmp_path):
        a = tmp_path / "a.md"
        a.write_text("# A", encoding="utf-8")
        conn = LocalFileConnector()
        conn.authenticate({
            "paths": [str(a)],
            "pattern": "*.md",  # non-recursive pattern — no '**'
            "recursive": "false",
            "follow_links": "true",
            "exclude": ["**/node_modules/**"],
        })
        assert conn._pattern == "*.md"
        assert conn._recursive is False
        assert conn._follow_links is True
        assert conn._exclude == ["**/node_modules/**"]

    def test_authenticate_rejects_glob_starstar_when_not_recursive(self, tmp_path):
        a = tmp_path / "a.md"
        a.write_text("# A", encoding="utf-8")
        conn = LocalFileConnector()
        with pytest.raises(ValueError, match="contains '\\*\\*'"):
            conn.authenticate({
                "paths": [str(a)],
                "pattern": "**/*.md",
                "recursive": "false",
            })


class TestLocalFileConnectorFetch:
    def _make_dir(self, tmp_path: Path) -> Path:
        root = tmp_path / "docs"
        root.mkdir()
        (root / "first.md").write_text(
            "# First Doc\n\nWelcome.\n\n## Section A\n\nbody A. See [docs](https://example.com).\n",
            encoding="utf-8",
        )
        (root / "second.md").write_text(
            "# Second Doc\n\n## Cross\n\nLink to [other](./first.md#section-a).\n",
            encoding="utf-8",
        )
        sub = root / "skip"
        sub.mkdir()
        (sub / "ignored.md").write_text("# ignored", encoding="utf-8")
        return root

    def test_fetch_markdown_directory(self, tmp_path):
        root = self._make_dir(tmp_path)
        conn = LocalFileConnector()
        conn.authenticate({"paths": [str(root)]})
        data = conn.fetch()
        assert isinstance(data, NormalizedData)
        # 2 real docs + 1 ignored doc (in subdir) + 1 URL stub.
        doc_names = {d["name"] for d in data.entities["Document"]}
        assert any(name.endswith("/first.md") for name in doc_names)
        assert any(name.endswith("/second.md") for name in doc_names)
        assert "https://example.com" in doc_names
        # Section URIs reflect the doc URI + slug.
        section_names = {s["name"] for s in data.entities["Section"]}
        assert any(name.endswith("/first.md#first-doc") for name in section_names)
        assert any(name.endswith("/first.md#first-doc/section-a") for name in section_names)
        # Cross-doc anchor link resolved.
        first_uri = next(
            d["name"] for d in data.entities["Document"]
            if d["name"].endswith("/first.md")
        )
        expected_target = f"{first_uri}#section-a"
        assert any(
            r["type"] == "LINKS_TO" and r["target_name"] == expected_target
            for r in data.relationships
        )

    def test_fetch_with_exclude_pattern(self, tmp_path):
        root = self._make_dir(tmp_path)
        conn = LocalFileConnector()
        conn.authenticate({
            "paths": [str(root)],
            "exclude": ["**/skip/**"],
        })
        data = conn.fetch()
        # The 'ignored.md' file is in skip/ and must be excluded.
        names = {d["name"] for d in data.entities["Document"]}
        assert not any(name.endswith("/skip/ignored.md") for name in names)

    def test_fetch_non_recursive(self, tmp_path):
        root = self._make_dir(tmp_path)
        conn = LocalFileConnector()
        conn.authenticate({
            "paths": [str(root)],
            "pattern": "*.md",
            "recursive": "false",
        })
        data = conn.fetch()
        names = {d["name"] for d in data.entities["Document"]}
        # Subdirectory should NOT have been walked.
        assert not any(name.endswith("/skip/ignored.md") for name in names)
        # Top-level files should still be present.
        assert any(name.endswith("/first.md") for name in names)

    def test_fetch_single_file_path(self, tmp_path):
        f = tmp_path / "only.md"
        f.write_text("# Solo\nbody\n", encoding="utf-8")
        conn = LocalFileConnector()
        conn.authenticate({"paths": [str(f)]})
        data = conn.fetch()
        # Exactly one Document plus its section.
        assert len(data.entities.get("Document", [])) == 1
        assert data.entities["Document"][0]["title"] == "Solo"

    def test_fetch_unsupported_extension_silently_skipped(self, tmp_path):
        root = tmp_path / "mixed"
        root.mkdir()
        (root / "good.md").write_text("# G", encoding="utf-8")
        (root / "bad.xyz").write_text("nope", encoding="utf-8")
        conn = LocalFileConnector()
        conn.authenticate({"paths": [str(root)]})
        data = conn.fetch()
        names = {d["name"] for d in data.entities["Document"]}
        assert any(name.endswith("/good.md") for name in names)
        assert not any(name.endswith("/bad.xyz") for name in names)

    def test_fetch_sorted_file_order(self, tmp_path):
        root = tmp_path / "ordered"
        root.mkdir()
        # Create files in non-alphabetical order — discovery must still
        # produce them in sorted-by-URI order.
        for stem in ("zeta", "alpha", "mid"):
            (root / f"{stem}.md").write_text(f"# {stem}\n", encoding="utf-8")
        conn = LocalFileConnector()
        conn.authenticate({"paths": [str(root)]})
        files = list(conn._discover_files())
        assert [f.name for f in files] == ["alpha.md", "mid.md", "zeta.md"]

    def test_fetch_idempotent(self, tmp_path):
        root = self._make_dir(tmp_path)
        conn1 = LocalFileConnector()
        conn1.authenticate({"paths": [str(root)]})
        conn2 = LocalFileConnector()
        conn2.authenticate({"paths": [str(root)]})
        a = conn1.fetch()
        b = conn2.fetch()
        # Strip loadedAt before structural comparison.
        for data in (a, b):
            for entity_list in data.entities.values():
                for entity in entity_list:
                    entity.pop("loadedAt", None)
        assert a.entities == b.entities
        # Relationships should match in count + content (order may vary if
        # we ever parallelise; here we just check membership).
        assert sorted(map(repr, a.relationships)) == sorted(map(repr, b.relationships))

    def test_uri_posix_form(self, tmp_path):
        f = tmp_path / "sub" / "doc.md"
        f.parent.mkdir()
        f.write_text("# X", encoding="utf-8")
        conn = LocalFileConnector()
        conn.authenticate({"paths": [str(f)]})
        data = conn.fetch()
        for d in data.entities["Document"]:
            if d["name"].startswith("http"):
                continue
            assert "\\" not in d["name"]

    def test_documents_list_empty(self, tmp_path):
        # Spec §5: connector must NOT populate the `documents=[]` list.
        f = tmp_path / "x.md"
        f.write_text("# T\nbody\n", encoding="utf-8")
        conn = LocalFileConnector()
        conn.authenticate({"paths": [str(f)]})
        data = conn.fetch()
        assert data.documents == []

    def test_merge_with_other_connector_results(self, tmp_path):
        # Verify outputs play well with merge_connector_results()
        a = tmp_path / "a.md"
        b = tmp_path / "b.md"
        a.write_text("# A\nbody\n", encoding="utf-8")
        b.write_text("# B\nbody\n", encoding="utf-8")
        c1 = LocalFileConnector()
        c1.authenticate({"paths": [str(a)]})
        c2 = LocalFileConnector()
        c2.authenticate({"paths": [str(b)]})
        merged = merge_connector_results([c1.fetch(), c2.fetch()])
        # Both documents should be present.
        names = {d["name"] for d in merged.entities["Document"]}
        assert any(name.endswith("/a.md") for name in names)
        assert any(name.endswith("/b.md") for name in names)


# ---------------------------------------------------------------------------
# File metadata: parse_file() enrichment (stat-derived fields)
# ---------------------------------------------------------------------------


class TestFileMetadataFromStat:
    """parse_file() populates stat-derived fields on ParsedDocument."""

    def test_file_extension_set(self, tmp_path):
        f = tmp_path / "guide.md"
        f.write_text("# Title\n\nbody\n")
        doc = parse_file(f)
        assert doc.file_extension == "md"

    def test_file_size_positive(self, tmp_path):
        f = tmp_path / "guide.md"
        content = "# Title\n\nbody\n"
        f.write_text(content)
        doc = parse_file(f)
        assert doc.file_size == f.stat().st_size
        assert doc.file_size > 0

    def test_created_at_is_datetime(self, tmp_path):
        from datetime import datetime, timezone
        f = tmp_path / "guide.md"
        f.write_text("# Title\n")
        doc = parse_file(f)
        assert isinstance(doc.created_at, datetime)
        assert doc.created_at.tzinfo is not None

    def test_modified_at_is_datetime(self, tmp_path):
        from datetime import datetime, timezone
        f = tmp_path / "guide.md"
        f.write_text("# Title\n")
        doc = parse_file(f)
        assert isinstance(doc.modified_at, datetime)
        assert doc.modified_at.tzinfo is not None

    def test_no_extension_gives_none(self, tmp_path):
        # Files without a supported extension raise ValueError; this tests
        # that the extension derivation itself returns None for bare stems.
        from create_context_graph.connectors._local_file.parser import ParsedDocument
        doc = ParsedDocument(uri="/tmp/README", title="README")
        assert doc.file_extension is None

    def test_html_extension(self, tmp_path):
        pytest.importorskip("bs4")
        f = tmp_path / "page.html"
        f.write_text("<html><body><h1>Title</h1></body></html>")
        doc = parse_file(f)
        assert doc.file_extension == "html"


# ---------------------------------------------------------------------------
# Per-format parser: author / language extraction
# ---------------------------------------------------------------------------


class TestMarkdownParserMetadata:
    def test_frontmatter_author(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("---\nauthor: Alice\n---\n# Title\n\nbody\n")
        doc = parse_file(f)
        assert doc.author == "Alice"

    def test_frontmatter_language(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("---\nlang: de\n---\n# Title\n")
        doc = parse_file(f)
        assert doc.language == "de"

    def test_frontmatter_language_key_alias(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("---\nlanguage: fr\n---\n# Title\n")
        doc = parse_file(f)
        assert doc.language == "fr"

    def test_frontmatter_authors_list(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("---\nauthors:\n  - Alice\n  - Bob\n---\n# Title\n")
        doc = parse_file(f)
        assert doc.author == "Alice, Bob"

    def test_no_frontmatter_gives_none(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("# Title\n\nno frontmatter\n")
        doc = parse_file(f)
        assert doc.author is None
        assert doc.language is None

    def test_empty_frontmatter_gives_none(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("---\n---\n# Title\n")
        doc = parse_file(f)
        assert doc.author is None
        assert doc.language is None


class TestHTMLParserMetadata:
    def test_meta_author(self, tmp_path):
        pytest.importorskip("bs4")
        f = tmp_path / "page.html"
        f.write_text(
            '<html><head><meta name="author" content="Bob"/></head>'
            '<body><h1>Title</h1></body></html>'
        )
        doc = parse_file(f)
        assert doc.author == "Bob"

    def test_html_lang_attribute(self, tmp_path):
        pytest.importorskip("bs4")
        f = tmp_path / "page.html"
        f.write_text('<html lang="en-US"><body><h1>Title</h1></body></html>')
        doc = parse_file(f)
        assert doc.language == "en-US"

    def test_meta_content_language_fallback(self, tmp_path):
        pytest.importorskip("bs4")
        f = tmp_path / "page.html"
        f.write_text(
            '<html><head><meta http-equiv="content-language" content="fr"/></head>'
            '<body><h1>Title</h1></body></html>'
        )
        doc = parse_file(f)
        assert doc.language == "fr"

    def test_no_metadata_gives_none(self, tmp_path):
        pytest.importorskip("bs4")
        f = tmp_path / "page.html"
        f.write_text("<html><body><h1>Title</h1><p>body</p></body></html>")
        doc = parse_file(f)
        assert doc.author is None
        assert doc.language is None


class TestAsciiDocParserMetadata:
    def test_author_attribute(self, tmp_path):
        f = tmp_path / "doc.adoc"
        f.write_text("= Title\n:author: Carol\n\nBody text.\n")
        doc = parse_file(f)
        assert doc.author == "Carol"

    def test_lang_attribute(self, tmp_path):
        f = tmp_path / "doc.adoc"
        f.write_text("= Title\n:lang: ja\n\nBody text.\n")
        doc = parse_file(f)
        assert doc.language == "ja"

    def test_language_attribute_alias(self, tmp_path):
        f = tmp_path / "doc.adoc"
        f.write_text("= Title\n:language: es\n\nBody text.\n")
        doc = parse_file(f)
        assert doc.language == "es"

    def test_no_attributes_gives_none(self, tmp_path):
        f = tmp_path / "doc.adoc"
        f.write_text("= Title\n\nJust a body.\n")
        doc = parse_file(f)
        assert doc.author is None
        assert doc.language is None


# ---------------------------------------------------------------------------
# Mapper: new metadata fields reach the entity dict and sections get
# fileExtension
# ---------------------------------------------------------------------------


class TestDocumentMapperMetadata:
    """Verify all new metadata properties flow from ParsedDocument → entity dict."""

    def _make_doc(self, **kwargs) -> ParsedDocument:
        from datetime import datetime, timezone
        defaults = dict(
            uri="/tmp/test/guide.md",
            title="Guide",
            preamble="intro",
            sections=[],
            links=[],
            source_type="LOCAL_FILE",
            file_extension="md",
            file_size=1234,
            created_at=datetime(2025, 1, 15, 10, 0, 0, tzinfo=timezone.utc),
            modified_at=datetime(2025, 6, 1, 12, 0, 0, tzinfo=timezone.utc),
            author="Alice",
            language="en",
            page_count=None,
        )
        defaults.update(kwargs)
        return ParsedDocument(**defaults)

    def test_file_metadata_on_document_entity(self):
        from create_context_graph.connectors._local_file.mapper import DocumentMapper
        mapper = DocumentMapper()
        mapper.add(self._make_doc())
        data = mapper.build()
        entity = data.entities["Document"][0]
        assert entity["fileExtension"] == "md"
        assert entity["fileSize"] == 1234
        assert entity["author"] == "Alice"
        assert entity["language"] == "en"

    def test_loadedAt_is_datetime_not_string(self):
        from datetime import datetime
        from create_context_graph.connectors._local_file.mapper import DocumentMapper
        mapper = DocumentMapper()
        mapper.add(self._make_doc())
        data = mapper.build()
        entity = data.entities["Document"][0]
        assert isinstance(entity["loadedAt"], datetime), (
            "loadedAt must be a native datetime for Neo4j ZonedDateTime storage"
        )

    def test_created_at_and_modified_at_stored(self):
        from datetime import datetime, timezone
        from create_context_graph.connectors._local_file.mapper import DocumentMapper
        mapper = DocumentMapper()
        doc = self._make_doc()
        mapper.add(doc)
        data = mapper.build()
        entity = data.entities["Document"][0]
        assert entity["createdAt"] == doc.created_at
        assert entity["modifiedAt"] == doc.modified_at

    def test_page_count_stored_when_present(self):
        from create_context_graph.connectors._local_file.mapper import DocumentMapper
        mapper = DocumentMapper()
        mapper.add(self._make_doc(page_count=42))
        data = mapper.build()
        entity = data.entities["Document"][0]
        assert entity["pageCount"] == 42

    def test_page_count_none_when_absent(self):
        from create_context_graph.connectors._local_file.mapper import DocumentMapper
        mapper = DocumentMapper()
        mapper.add(self._make_doc(page_count=None))
        data = mapper.build()
        entity = data.entities["Document"][0]
        assert entity["pageCount"] is None

    def test_section_gets_file_extension(self):
        from create_context_graph.connectors._local_file.mapper import DocumentMapper
        from create_context_graph.connectors._local_file.parser import ParsedSection
        doc = self._make_doc(sections=[
            ParsedSection(title="Intro", level=1, body="hello", subsections=[], links=[])
        ])
        mapper = DocumentMapper()
        mapper.add(doc)
        data = mapper.build()
        sections = data.entities.get("Section", [])
        assert sections, "expected at least one Section entity"
        assert sections[0]["fileExtension"] == "md"

    def test_section_file_extension_from_pdf_uri(self):
        from create_context_graph.connectors._local_file.mapper import DocumentMapper
        from create_context_graph.connectors._local_file.parser import ParsedSection
        doc = self._make_doc(
            uri="/tmp/report.pdf",
            file_extension="pdf",
            sections=[
                ParsedSection(title="Ch1", level=1, body="text", subsections=[], links=[])
            ]
        )
        mapper = DocumentMapper()
        mapper.add(doc)
        data = mapper.build()
        section = data.entities["Section"][0]
        assert section["fileExtension"] == "pdf"

    def test_stub_upgrade_preserves_metadata(self):
        """When a stub Document is later ingested, metadata is filled in correctly."""
        from create_context_graph.connectors._local_file.mapper import DocumentMapper
        from create_context_graph.connectors._local_file.parser import ParsedSection
        # First doc links to /tmp/linked.md — creates a stub.
        linker = self._make_doc(
            uri="/tmp/source.md",
            sections=[
                ParsedSection(title="S1", level=1, body="See [other](linked.md)", subsections=[], links=["linked.md"])
            ]
        )
        mapper = DocumentMapper()
        mapper.add(linker)

        # Now ingest the linked doc as a real document.
        linked = self._make_doc(
            uri="/tmp/linked.md",
            title="Linked",
            author="Bob",
            language="fr",
            file_extension="md",
            file_size=500,
        )
        mapper.add(linked)
        data = mapper.build()

        docs = {d["name"]: d for d in data.entities["Document"]}
        assert "/tmp/linked.md" in docs
        linked_entity = docs["/tmp/linked.md"]
        assert linked_entity.get("title") == "Linked"
        assert linked_entity.get("author") == "Bob"
        assert linked_entity.get("language") == "fr"
